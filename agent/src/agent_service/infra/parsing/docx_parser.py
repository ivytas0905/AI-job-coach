"""DOCX Parser Implementation"""
import io
from docx import Document


class DOCXParser:
    """Extract text from DOCX files"""

    @staticmethod
    def extract_text(file_content: bytes) -> str:
        """
        Extract text from DOCX file

        Args:
            file_content: DOCX file bytes

        Returns:
            Extracted text content
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n".join(text_content)

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
